from io import BytesIO
from flask import send_file
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib import utils
from docx import Document
from docx.shared import Inches
import json
import base64

def wrap_text(text, max_width, font_name='Helvetica', font_size=12):
    """
    Wraps the text to fit within the specified max width.
    """
    from reportlab.pdfgen.canvas import Canvas
    temp_canvas = Canvas(None)
    temp_canvas.setFont(font_name, font_size)
    lines = text.splitlines()
    wrapped_lines = []
    for line in lines:
        if temp_canvas.stringWidth(line) > max_width:
            words = line.split(' ')
            current_line = ''
            for word in words:
                if temp_canvas.stringWidth(current_line + ' ' + word) <= max_width:
                    current_line += ' ' + word
                else:
                    wrapped_lines.append(current_line.strip())
                    current_line = word
            wrapped_lines.append(current_line.strip())
        else:
            wrapped_lines.append(line)
    return wrapped_lines

def export_pdf(text, date, kg_data, entities):
    """Generates a PDF file with the given text, date, kg data, and entities."""
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    margin = inch
    text_width = width - 2 * margin
    textobject = p.beginText(margin, height - margin)
    textobject.setFont("Helvetica", 12)
    wrapped_lines = wrap_text(f"Summary Date: {date}\nSummary Content:\n{text}", text_width)
    for line in wrapped_lines:
        textobject.textLine(line)
    p.drawText(textobject)

    if kg_data:
        p.showPage()
        textobject = p.beginText(margin, height - margin)
        textobject.setFont("Helvetica", 12)
        textobject.textLine("Knowledge Graph:")
        p.drawText(textobject)

        for image in json.loads(kg_data):
            img_data = base64.b64decode(image)
            img_buffer = BytesIO(img_data)
            img = utils.ImageReader(img_buffer)
            img_width, img_height = img.getSize()
            aspect = img_height / float(img_width)

            if img_width > width - 2 * margin:
                img_width = width - 2 * margin
                img_height = img_width * aspect

            if img_height > height - 2 * margin:
                img_height = height - 2 * margin
                img_width = img_height / aspect

            p.drawImage(img, margin, height - 2 * margin - img_height, width=img_width, height=img_height, preserveAspectRatio=True, mask='auto')
            p.showPage()  # Add a new page for each image

    if entities:
        p.showPage()
        textobject = p.beginText(margin, height - margin)
        textobject.setFont("Helvetica", 12)
        textobject.textLine("Entities:")
        p.drawText(textobject)
        textobject = p.beginText(margin, height - 2 * margin)
        for entity_type, entity_list in json.loads(entities).items():
            textobject.textLine(f"{entity_type}:")
            for entity in entity_list:
                textobject.textLine(f" - {entity}")
        p.drawText(textobject)

    p.showPage()
    p.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='summary.pdf', mimetype='application/pdf')

def export_word(text, date, kg_data, entities):
    """Generates a Word document with the given text, date, kg data, and entities."""
    buffer = BytesIO()
    doc = Document()
    doc.add_heading('Summary Details', level=1)
    doc.add_paragraph(f"Summary Date: {date}")
    doc.add_paragraph("Summary Content:")
    doc.add_paragraph(text)

    if kg_data:
        doc.add_heading('Knowledge Graph:', level=2)
        for image in json.loads(kg_data):
            img_data = base64.b64decode(image)
            img_buffer = BytesIO(img_data)
            doc.add_picture(img_buffer, width=Inches(4))

    if entities:
        doc.add_heading('Entities:', level=2)
        for entity_type, entity_list in json.loads(entities).items():
            doc.add_paragraph(f"{entity_type}:")
            for entity in entity_list:
                doc.add_paragraph(f" - {entity}", style='ListBullet')

    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='summary.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def export_txt(text, date, kg_data, entities):
    """Generates a text file with the given text, date, kg data, and entities."""
    buffer = BytesIO()
    content = f"Summary Date: {date}\nSummary Content:\n{text}\n"

    if entities:
        content += "\nEntities:\n"
        for entity_type, entity_list in json.loads(entities).items():
            content += f"{entity_type}:\n"
            for entity in entity_list:
                content += f" - {entity}\n"

    buffer.write(content.encode('utf-8'))
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='summary.txt', mimetype='text/plain')